import os
import re
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

from bedrock_client import (
    generate_about_customer,
    generate_project_objectives,
    generate_current_landscape,
    generate_key_highlights,
    generate_dr_requirements,
    generate_proposed_drs_solution,
    generate_tco_assumptions,
    generate_high_level_scope,
    generate_scope_of_work,
    generate_project_deliverables,
    generate_acceptance_criteria,
    generate_customer_obligations,
    generate_customer_dependencies,
    generate_assumptions,
    generate_exclusions,
)
from pricing_client import build_pricing_estimate

# ── helpers ──────────────────────────────────────────────────────────────────

ORANGE = RGBColor(0xFF, 0x6B, 0x00)   # Operisoft brand orange
DARK   = RGBColor(0x1F, 0x27, 0x63)   # dark navy

# ── Document font constants ──
FONT_NAME = "Calibri"
FONT_HEADING1_SIZE = Pt(16)
FONT_HEADING2_SIZE = Pt(13)
FONT_BODY_SIZE = Pt(10.5)
FONT_SMALL_SIZE = Pt(9)


def _set_document_defaults(doc):
    """Set default font to Calibri for the entire document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_BODY_SIZE
    # Set paragraph spacing defaults
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def _add_header_footer(doc, customer_name):
    """Add footer to the document (no header per template design)."""
    for section in doc.sections:
        # ── No Header — keep it empty ──
        header = section.header
        header.is_linked_to_previous = False
        # Clear any default header content
        for p in header.paragraphs:
            p.clear()

        # ── Footer ──
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add top border line (horizontal rule)
        pPr_f = footer_para._p.get_or_add_pPr()
        pBdr_f = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), '000000')
        pBdr_f.append(top)
        pPr_f.append(pBdr_f)

        # Line 1: "Operisoft Technologies Pvt Ltd" — right aligned
        run_f = footer_para.add_run("Operisoft Technologies Pvt Ltd")
        run_f.font.name = "Calibri (Body)"
        run_f.font.size = Pt(11)
        run_f.font.bold = True
        run_f.font.color.rgb = DARK

        # Line 2: CIN / GSTIN / website / email
        footer_para2 = footer.add_paragraph()
        footer_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_cin = footer_para2.add_run("CIN: U72900MH2022PTC384302| GSTIN: ")
        run_cin.font.name = "Calibri (Body)"
        run_cin.font.size = Pt(11)
        run_cin.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run_gstin = footer_para2.add_run("27AADCO8045M1ZZ")
        run_gstin.font.name = "Calibri (Body)"
        run_gstin.font.size = Pt(11)
        run_gstin.font.bold = True
        run_gstin.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run_rest = footer_para2.add_run(" | www.operisoft.com | E: info@operisoft.com |")
        run_rest.font.name = "Calibri (Body)"
        run_rest.font.size = Pt(11)
        run_rest.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Line 3: Page number — right aligned
        footer_para3 = footer.add_paragraph()
        footer_para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_page = footer_para3.add_run("Page | ")
        run_page.font.name = "Calibri (Body)"
        run_page.font.size = Pt(11)
        run_page.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Add page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_page._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        run_page._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_page._r.append(fldChar2)


def _toc_entry(doc, text, level=0):
    """Add a single TOC entry with dot leader tab stop matching Word's native TOC style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

    if level == 0:
        p.paragraph_format.left_indent = Inches(0)
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Calibri (Body)"
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
    else:
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(text)
        run.bold = False
        run.font.name = "Calibri (Body)"
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        run.font.small_caps = True

    # Add right-aligned tab stop with dot leader for page number
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    # Right tab at ~16cm (page width minus margins)
    tab.set(qn('w:pos'), '9072')
    tabs.append(tab)
    pPr.append(tabs)


def _build_toc(doc, project_type, include_post_deploy, include_landing_zone,
               include_control_tower, include_lz_arch, include_paloalto,
               include_mgn, include_testing_monitoring, include_monitoring, include_dr):
    """Build a dynamic Table of Contents based on which sections are included."""

    # TOC Title — "TABLE OF CONTENTS" with bottom border
    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.space_after = Pt(12)
    run_t = toc_title.add_run("Table of Contents")
    run_t.bold = True
    run_t.font.name = "Calibri (Body)"
    run_t.font.size = Pt(24)
    run_t.font.color.rgb = DARK
    run_t.font.small_caps = True
    # Bottom border line under the title
    pPr = toc_title._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Build entries dynamically
    _toc_entry(doc, "Document History", level=0)

    _toc_entry(doc, "1.  Executive Summary", level=0)
    _toc_entry(doc, "1.1   Confidentiality Notice", level=1)
    _toc_entry(doc, "1.2   About Operisoft", level=1)
    _toc_entry(doc, "1.3   About Customer", level=1)

    _toc_entry(doc, "2.  Project Details", level=0)
    _toc_entry(doc, "2.1   Project Objectives and Overview", level=1)
    _toc_entry(doc, "2.2   Understanding of Current Landscape", level=1)

    # Section 3 title depends on project type
    if project_type.strip().upper() == "POC":
        _toc_entry(doc, "3.  Proposed Solution - POC", level=0)
    else:
        _toc_entry(doc, "3.  Proposed Solution", level=0)

    _toc_entry(doc, "3.1   High Level Scope", level=1)
    _toc_entry(doc, "3.2   Scope of Work", level=1)
    _toc_entry(doc, "3.3   Key Highlights of the Solution", level=1)
    _toc_entry(doc, "3.4   Proposed Solution Diagram", level=1)
    _toc_entry(doc, "3.5   Project Deliverables", level=1)

    if project_type.strip().upper() == "POC":
        _toc_entry(doc, "3.6   POC Acceptance Criteria", level=1)
    else:
        _toc_entry(doc, "3.6   Project Acceptance Criteria", level=1)

    if include_landing_zone:
        _toc_entry(doc, "3.7   Implementation of Landing Zone", level=1)
    if include_control_tower:
        _toc_entry(doc, "3.8   Configuration of Control Tower Setup", level=1)
    if include_lz_arch:
        _toc_entry(doc, "3.9   AWS Landing Zone Architecture", level=1)
    if include_paloalto:
        _toc_entry(doc, "3.10  PaloAlto Next Generation Firewall", level=1)

    if include_post_deploy:
        _toc_entry(doc, "4.  Post-Deployment Testing and Acceptance", level=0)

    if include_mgn or include_testing_monitoring:
        _toc_entry(doc, "4.  Migration Approach", level=0)
        if include_mgn:
            _toc_entry(doc, "4.1   Migration of Data Using MGN", level=1)
        if include_testing_monitoring:
            _toc_entry(doc, "4.2   Testing and Monitoring", level=1)

    if include_monitoring:
        _toc_entry(doc, "5.  Monitoring AWS Infrastructure", level=0)

    if include_dr:
        _toc_entry(doc, "6.  Proposed DR Approach", level=0)
        _toc_entry(doc, "6.1   Requirement of DR", level=1)
        _toc_entry(doc, "6.2   AWS Elastic Disaster Recovery for Replication", level=1)
        _toc_entry(doc, "6.3   AWS DRS Architecture Design", level=1)
        _toc_entry(doc, "6.4   Proposed DRS Solution", level=1)

    _toc_entry(doc, "5.  Cost Estimate/Commercial Terms", level=0)
    _toc_entry(doc, "5.1   AWS Total Cost of Ownership [TCO]", level=1)
    _toc_entry(doc, "5.2   Operisoft Service Charges", level=1)
    _toc_entry(doc, "5.3   Estimated Timeline", level=1)
    _toc_entry(doc, "5.4   Assumptions Taken for Cost Calculation", level=1)

    _toc_entry(doc, "6.  AWS Security Best Practices (Suggested)", level=0)
    _toc_entry(doc, "6.1   AWS Identity and Access Management (IAM)", level=1)
    _toc_entry(doc, "6.2   AWS IAM Access Analyzer", level=1)
    _toc_entry(doc, "6.3   Detective Controls", level=1)
    _toc_entry(doc, "6.4   AWS Detective", level=1)
    _toc_entry(doc, "6.5   AWS Security Hub", level=1)

    _toc_entry(doc, "7.  Customer Obligations and Engagement Terms", level=0)

    _toc_entry(doc, "8.  Customer Dependencies", level=0)

    _toc_entry(doc, "9.  Assumptions", level=0)

    _toc_entry(doc, "10. Exclusions", level=0)

    _toc_entry(doc, "11. Data Ownership and Customer Offboarding", level=0)
    _toc_entry(doc, "11.1  Data Ownership", level=1)
    _toc_entry(doc, "11.2  Offboarding and Data Handover", level=1)
    _toc_entry(doc, "11.3  Access Deprovisioning", level=1)
    _toc_entry(doc, "11.4  Transition Support (Optional)", level=1)

    _toc_entry(doc, "12. Commercial Terms and Conditions", level=0)


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading %d" % level] if ("Heading %d" % level) in [s.name for s in doc.styles] else doc.styles["Normal"]
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri (Body)"
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.small_caps = True
        # Add bottom border (horizontal line) below main headings
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        # Sub-heading: size 14, Blue Accent 1 Darker 25%, with left indent
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)  # Blue, Accent 1, Darker 25%
        p.paragraph_format.left_indent = Cm(0.7)
    return p

def _para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri (Body)"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def _bullet(doc, text, size=11):
    """Add a proper Word bullet point using the List Bullet style."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri (Body)"
    run.font.size = Pt(size)
    return p

def _add_bullets_from_text(doc, text):
    """Parse bullet text (lines starting with ➢, •, Ø, or -) and add as bullet paras."""
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        line = line.lstrip("•-·Ø➢").strip()
        if line:
            _bullet(doc, line)


def _render_ai_content(doc, text):
    """Render AI-generated content intelligently — bullets as real Word bullets, rest as paragraphs.
    Lines starting with ➢, Ø, •, or - become Word bullet points.
    Lines starting with · (sub-bullet) become indented Word bullets.
    Other lines become normal paragraphs.
    """
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Detect sub-bullet (·)
        if line.startswith("·"):
            clean = line.lstrip("·").strip()
            if clean:
                p = doc.add_paragraph(style="List Bullet 2") if "List Bullet 2" in [s.name for s in doc.styles] else doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(2.2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(clean)
                run.font.name = "Calibri (Body)"
                run.font.size = Pt(11)
        # Detect top-level bullet (➢, Ø, •, -)
        elif line[0] in "➢Ø•-" or line.startswith("- "):
            clean = line.lstrip("➢Ø•-").strip()
            if clean:
                _bullet(doc, clean)
        else:
            # Regular paragraph text (section labels like "Notes:", "Current Scenario", etc.)
            _para(doc, line)

def _section_divider(doc):
    doc.add_paragraph()

def _two_col_table(doc, rows_data, header=None):
    """Create a simple 2-column table."""
    cols = 2
    table = doc.add_table(rows=len(rows_data) + (1 if header else 0), cols=cols)
    table.style = "Table Grid"
    if header:
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(hdr_cells[i], "1F2763")
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for idx, (k, v) in enumerate(rows_data):
        row = table.rows[idx + (1 if header else 0)]
        row.cells[0].text = k
        row.cells[1].text = v
    return table


# ── static content blocks ──────────────────────────────────────────────────

CONFIDENTIALITY_TEXT = (
    "The information contained in this document has been prepared to be used in the context "
    "of this project. It should not be used as a model or precedent in any situation outside "
    "of this project. This document must not be copied or reproduced by any means without the "
    "authorization of the parties involved. A good effort has gone into the preparation of this "
    "document to ensure that the information presented is correct at the time of printing. The "
    "parties involved in this project assume no responsibility for any errors that may arise in "
    "the application of this information in a context other than the project for which it was prepared."
)

ABOUT_OPERISOFT = (
    "Operisoft Technologies Private Limited is a leading software consulting and services provider "
    "in digital transformation and Cloud Technologies, headquartered in Mumbai with offices in Pune, "
    "Bangalore, Hyderabad and Ahmedabad. We specialize in cloud technologies, including Amazon Web "
    "Services (AWS), providing expert support for complex solutions and the latest cloud innovations. "
    "With a team of over 55+ skilled professionals across India, we deliver best-in-class solutions "
    "tailored to customer needs.\n\n"
    "Operisoft supports clients across diverse sectors—including manufacturing, pharmaceuticals, "
    "healthcare, automotive, plant and machinery, chemical and process industries, e-commerce and "
    "retail, media and entertainment, as well as education and non-profit organizations—by leveraging "
    "cloud technologies like AWS and SaaS platforms. We proudly serve over 200 customers across these industries.\n\n"
    "With over 5+ years of experience as a leading AWS Cloud Partner, Operisoft excels in guiding "
    "businesses through cloud migration and modernization. Our team boasts the highest credentials, "
    "including 3 major AWS competencies & service validations, and certified AWS specialists, with "
    "solution architects holding 20+ active AWS certifications. As an AWS Advanced Tier Services Partner, "
    "we showcase excellence in AWS Managed Services, Data & Analytics, Migration, and DevOps. We were "
    "honoured as the Rising Star Partner of the Year 2023 in India."
)

TESTING_MONITORING_TEXT = [
    "Testing of the application will be done by client after the deployment.",
    "Once the Testing of the application is successful, the DNS records will be updated for going live.",
]

MONITORING_INFRA_BULLETS = [
    "Detailed monitoring will be done for all servers using AWS CloudWatch.",
    "For monitoring of OS level metrics like memory and disk utilization, custom CloudWatch metrics to be used.",
    "We will configure CloudWatch agent for custom CloudWatch monitoring.",
    "CloudWatch Alerts will be configured if resource utilization on any of the servers crosses a specific threshold.",
    "Monitoring parameters and thresholds will be configured as per the customer KPIs.",
    "Alerts will be configured for the status of the server for any issue on the server and server is not ready status.",
]

EDR_TEXT = (
    "AWS EDR (Elastic Disaster Recovery) minimizes downtime and data loss with fast, reliable recovery "
    "of on-premises and cloud-based applications using affordable storage, minimal compute, and "
    "point-in-time recovery.\n\nDisaster Recovery utilizes block-level, Continuous Data Replication, "
    "which ensures that target machines are spun up in their most up-to-date state during a disaster "
    "or drill. Organizations can thereby achieve sub-second Recovery Point Objectives (RPOs)."
)

DRS_ARCH_BULLETS = [
    "Replication server status reporting",
    "Staging area resources automatically created and terminated",
    "Recovery instances launched with RTO of minutes and RPO of few Minutes",
    "Continuous block-level replication (compressed and encrypted)",
]

IAM_BULLETS = [
    "The root account will not be used for day-to-day activities and the credentials of the root account will remain with the customer team only.",
    "AWS console will be accessed through cross-account IAM roles. We recommend using SAML-based federated users to access the AWS account instead of IAM users.",
    "MFA will be enabled for root AWS accounts, and we suggest and strongly recommend using MFA for other IAM users too.",
    "AWS partner users will use cross-account IAM roles to access the AWS account and there will be no IAM user for the AWS partner team.",
    "A strong password policy will be in place which includes regular password rotation.",
    "Password length should not be less than 8 characters.",
    "Password should contain at least one character from a-z, A-Z, 0-9, and special characters.",
    "Old passwords cannot be repeated.",
    "A role-based access policy with the principle of least privilege will be configured for accessing the AWS console and other maintenance tasks.",
]

IAM_ANALYZER_BULLETS = [
    "AWS IAM access analyzer helps you to fine-grain the continued cycle of access to the least privileged.",
    "AWS IAM Access Analyzer helps you streamline permissions management throughout each step of the cycle.",
    "Access Analyzer will identify the resources shared with external principles by using logic-based reasoning to analyze resource-based policy in your AWS environment.",
    "Access analyzer can help you to validate your created policy against IAM best practices.",
    "Access analyzer will be able to analyze your AWS CloudTrail log to identify actions and services that have been used by an IAM entity.",
]

DETECTIVE_CONTROLS_BULLETS = [
    "CloudTrail Logs will be enabled for all users. The last 7 days' data can be monitored from AWS CloudTrail Console.",
    "VPC Flow logs will be enabled, and Guard Duty will be configured to monitor VPC Flow Logs.",
    "GuardDuty will generate findings whenever it detects unexpected and potentially malicious activity in your AWS environment.",
    "CloudTrail Logs will be stored in a private S3 bucket. Only permitted users will have access permission to that bucket.",
    "We recommend pushing access and other logs from all servers to CloudWatch so that logs can be analyzed.",
    "All the sys logs and application logs will be stored in S3 for 90 days.",
    "After 90 days, logs will be moved from S3 to Glacier.",
    "Logs in Glacier will be retained for another 90 days before they are permanently removed.",
    "Logs stored on S3 can be analyzed using AWS Athena.",
    "AWS config rules will be configured to monitor Config logs stored on the s3 bucket and used for dynamic compliance checking.",
    "CloudTrail Logs will be streamed to AWS CloudWatch and filters will be created in CloudWatch for modification in security groups, start-stop of EC2 instances, or use of root user APIs.",
]

AWS_DETECTIVE_BULLETS = [
    "Amazon Detective will enable you to analyze and visualize security data from your AWS CloudTrail logs, VPC Flow logs, and Amazon GuardDuty findings.",
    "Amazon Detective will be integrated with Amazon GuardDuty and AWS Security Hub as well as AWS partner security products.",
    "AWS Detective will provide an interactive unified view which will help you to visualize all the context and details in one place.",
]

SECURITY_HUB_BULLETS = [
    "AWS Security Hub will centralize and prioritize security findings from across AWS accounts, services, and supported third-party partners to help you analyze your security trends and identify the highest priority security issues.",
    "Security Hub will help you manage security across multiple AWS accounts within a region by configuring the multi-account hierarchy within AWS Security Hub.",
    "Security Hub will be deployed with CloudFormation Template.",
    "Security Hub can be integrated with AWS Detective, Guard Duty, and other security services.",
    "Creation and Configuration of security events alerts from AWS Detective.",
]







POST_DEPLOYMENT_BULLETS = [
    "Following deployment to any designated environment, the Client shall have a maximum of two (2) iterations or fourteen (14) calendar days, whichever occurs first to complete testing and provide written feedback.",
    "Operisoft will address in-scope issues reported within this period as part of the two iterations. Any requests beyond scope or submitted after this period will require a formal Change Order and may incur additional costs.",
    "Operisoft shall not be liable for any delays, rework, or impacts resulting from feedback, defects, or change requests that are submitted outside the defined timeline or that fall outside the agreed scope of work.",
    "If no feedback is received within the specified period, or all reported issues are resolved within the allowed iterations, the deployment shall be deemed accepted.",
]





COMMERCIAL_TERMS = """i. Offer Validity: Upto 30 days.

ii. Payment Term:
   a. Implementation: 100% advance unless explicitly mentioned.
   b. Managed Services/AWS Bills: 15 days after Invoice generation.

iii. Taxes: Extra @ 18% IGST/CGST+SGST in case of INR billing

iv. SAC/HSN:
   a. 998315 Cloud consumption & services
   b. 997331 Software

v. Billing Cycle: Monthly invoices are generated from 5th to 15th of the next month. And must be paid within 15 days to avoid stopping the services.

vi. Order to be placed on:
OPERISOFT TECHNOLOGIES PRIVATE LIMITED
Office No 301, Plot No 273, Zion Complex Building
Sector-10, Khaghar, Navi Mumbai – 410210"""


LANDING_ZONE_MANDATORY_GUARDRAILS = [
    "Disallow Changes to Encryption Configuration for AWS Control Tower Created Amazon S3 Buckets in Log Archive",
    "Disallow Changes to Logging Configuration for AWS Control Tower Created Amazon S3 Buckets in Log Archive",
    "Disallow Changes to Bucket Policy for AWS Control Tower Created Amazon S3 Buckets in Log Archive",
    "Disallow Changes to Lifecycle Configuration for AWS Control Tower Created Amazon S3 Buckets in Log Archive",
    "Disallow Changes to Amazon CloudWatch Logs Log Groups set up by AWS Control Tower",
    "Disallow Deletion of AWS Config Aggregation Authorizations Created by AWS Control Tower",
    "Disallow Deletion of Log Archive",
    "Detect Public Read Access Setting for Log Archive",
    "Detect Public Write Access Setting for Log Archive",
    "Disallow Configuration Changes to CloudTrail",
    "Integrate CloudTrail Events with Amazon CloudWatch Logs",
    "Enable CloudTrail in All Available Regions",
    "Enable Integrity Validation for CloudTrail Log File",
    "Disallow Changes to Amazon CloudWatch Set Up by AWS Control Tower",
    "Disallow Changes to Tags Created by AWS Control Tower for AWS Config Resources",
    "Disallow Configuration Changes to AWS Config",
    "Enable AWS Config in All Available Regions",
    "Disallow Changes to AWS Config Rules Set Up by AWS Control Tower",
    "Disallow Changes to AWS IAM Roles Set Up by AWS Control Tower and AWS CloudFormation",
    "Disallow Changes to AWS Lambda Functions Set Up by AWS Control Tower",
    "Disallow Changes to Amazon SNS Set Up by AWS Control Tower",
    "Disallow Changes to Amazon SNS Subscriptions Set Up by AWS Control Tower",
]

LANDING_ZONE_RECOMMENDED_GUARDRAILS = [
    "Disallow the Creation of Access Keys for the Root User",
    "Disallow Actions as a Root User",
    "Detect Whether Encryption is Enabled for Amazon EBS Volumes Attached to Amazon EC2 Instances",
    "Detect Whether Unrestricted Incoming TCP Traffic is Allowed",
    "Detect Whether Unrestricted Internet Connection Through SSH is Allowed",
    "Detect Whether MFA for the Root User is Enabled",
    "Detect Whether Public Read Access to Amazon S3 Buckets is Allowed",
    "Detect Whether Public Write Access to Amazon S3 Buckets is Allowed",
    "Detect Whether Amazon EBS Volumes are Attached to Amazon EC2 Instances",
    "Detect Whether Amazon EBS Optimization is Enabled for Amazon EC2 Instances",
    "Detect Whether Public Access to Amazon RDS Database Instances is Enabled",
    "Detect Whether Public Access to Amazon RDS Database Snapshots is Enabled",
    "Detect Whether Storage Encryption is Enabled for Amazon RDS Database Instances",
    "Enable Encryption at Rest for Log Archive",
    "Enable Access Logging for Log Archive",
    "Disallow Policy Changes to Log Archive",
    "Set a Retention Policy for the Log Archive",
    "Disallow Changes to Replication Configuration for Amazon S3 Buckets",
    "Disallow Delete Actions on CloudTrail S3 Buckets Without MFA",
    "Detect Whether MFA is Enabled for AWS IAM Users",
    "Detect Whether MFA is Enabled for AWS IAM Users of the AWS Console",
    "Detect Whether Versioning for Amazon S3 Buckets is Enabled",
]

PALOALTO_BULLETS = [
    "Application visibility for informed security decisions: The VM-Series provides application visibility across all ports.",
    "Segment/Whitelist applications for security and compliance: Using segmentation and whitelisting policies allows you to control applications communicating across different subnets.",
    "Prevent advanced attacks within allowed application flows: The VM-Series allows you to use Palo Alto Networks Threat Prevention, DNS Security, and WildFire® to apply application-specific policies.",
    "Control application access with user-based policies: Integration with Active Directory®, LDAP, and other user repositories.",
    "Policy consistency through centralized management: Panorama™ provides centralized network security management for your VM-Series firewalls.",
    "Cloud-native scalability and availability: In public cloud environments, we recommended using cloud services such as application gateways, load balancers, and automation.",
]

MGN_BULLETS = [
    "AWS partner team will set up the network infrastructure - VPN from Onprem/Other to AWS for secure replication.",
    "AWS partner team will create a separate VPC environment for isolation and security.",
    "Deployment of AWS MGN tool in Customers existing setup to initiate the replication of server data.",
    "Setup of Staging VPC and replication VM in AWS with EBS volume.",
    "AWS MGN is a highly automated lift-and-shift solution, which works by replicating (physical or virtual) and/or cloud servers into assigned AWS account.",
    "First time setup: Creating the replication template by initializing the AWS Application Migration Service.",
    "Adding source servers: Install the AWS Replication Agent on both Linux and Windows source servers.",
    "Configuring launch settings: Configure instructions that determine how a test or cutover instance will be launched.",
    "Launching a test instance: Test the migration of source servers to AWS prior to initiating a cutover.",
    "Launching a cutover instance: Once testing is finalized, the cutover will migrate source servers to AWS.",
]


# ── main document builder ─────────────────────────────────────────────────

def generate_sow_document(data: dict) -> str:
    customer_name         = data.get("customer_name", "Customer")
    mom_text              = data.get("mom_text", "")
    company_url           = data.get("company_url", "")
    project_type          = data.get("project_type", "POC")   # "POC" or "Production"
    doc_date              = data.get("doc_date", "")
    submitted_by          = data.get("submitted_by", "")
    client_logo_path      = data.get("client_logo_path", None)
    include_landing_zone  = data.get("include_landing_zone", False)
    include_control_tower = data.get("include_control_tower", False)
    include_lz_arch       = data.get("include_landing_zone_arch", False)
    include_paloalto      = data.get("include_paloalto", False)
    include_mgn           = data.get("include_mgn_migration", False)
    include_post_deploy   = data.get("include_post_deployment", False)
    include_testing_monitoring = data.get("include_testing_monitoring", False)
    include_monitoring    = data.get("include_monitoring", False)
    include_dr            = data.get("include_dr", False)

    # Extract company info from URL if provided
    company_about_us = ""
    if company_url and company_url.strip():
        from url_extractor import extract_company_info
        extracted = extract_company_info(company_url.strip())
        company_about_us = extracted.get("about_us", "")
        # Also use extracted company name if customer_name wasn't explicitly set
        if not customer_name or customer_name == "Customer":
            customer_name = extracted.get("company_name", customer_name)

    # Generate dynamic content via Bedrock (credentials loaded from .env)
    print(f"\n[SOW] ═══════════════════════════════════════════════════════════")
    print(f"[SOW] Starting SOW generation for: {customer_name}")
    print(f"[SOW] Project Type: {project_type}")
    print(f"[SOW] ═══════════════════════════════════════════════════════════")

    print(f"[SOW] [1/15] Generating 'About Customer'...")
    about_customer      = generate_about_customer(customer_name, mom_text, company_about_us)
    print(f"[SOW] [1/15] ✓ About Customer generated ({len(about_customer)} chars)")

    print(f"[SOW] [2/15] Generating 'Project Objectives'...")
    proj_objectives     = generate_project_objectives(customer_name, mom_text, project_type)
    print(f"[SOW] [2/15] ✓ Project Objectives generated ({len(proj_objectives)} chars)")

    print(f"[SOW] [3/15] Generating 'Current Landscape'...")
    current_landscape   = generate_current_landscape(customer_name, mom_text)
    print(f"[SOW] [3/15] ✓ Current Landscape generated ({len(current_landscape)} chars)")

    print(f"[SOW] [4/15] Generating 'Key Highlights'...")
    key_highlights      = generate_key_highlights(customer_name, mom_text)
    print(f"[SOW] [4/15] ✓ Key Highlights generated ({len(key_highlights)} chars)")

    print(f"[SOW] [5/15] Generating 'DR Requirements'...")
    dr_reqs             = generate_dr_requirements(customer_name, mom_text)
    print(f"[SOW] [5/15] ✓ DR Requirements generated")

    print(f"[SOW] [6/15] Generating 'Proposed DRS Solution'...")
    proposed_drs        = generate_proposed_drs_solution(customer_name, mom_text)
    print(f"[SOW] [6/15] ✓ Proposed DRS Solution generated ({len(proposed_drs)} chars)")

    print(f"[SOW] [7/15] Generating 'High Level Scope'...")
    high_level_scope    = generate_high_level_scope(customer_name, mom_text)
    print(f"[SOW] [7/15] ✓ High Level Scope generated ({len(high_level_scope)} chars)")

    print(f"[SOW] [8/15] Generating 'Scope of Work'...")
    scope_of_work       = generate_scope_of_work(customer_name, mom_text, project_type)
    print(f"[SOW] [8/15] ✓ Scope of Work generated ({len(scope_of_work)} chars)")

    print(f"[SOW] [9/15] Generating 'Project Deliverables'...")
    project_deliverables = generate_project_deliverables(customer_name, mom_text)
    print(f"[SOW] [9/15] ✓ Project Deliverables generated ({len(project_deliverables)} chars)")

    print(f"[SOW] [10/15] Generating 'Acceptance Criteria'...")
    acceptance_criteria = generate_acceptance_criteria(customer_name, mom_text, project_type)
    print(f"[SOW] [10/15] ✓ Acceptance Criteria generated ({len(acceptance_criteria)} chars)")

    print(f"[SOW] [11/15] Generating 'Customer Obligations'...")
    customer_obligations = generate_customer_obligations(customer_name, mom_text)
    print(f"[SOW] [11/15] ✓ Customer Obligations generated ({len(customer_obligations)} chars)")

    print(f"[SOW] [12/15] Generating 'Customer Dependencies'...")
    customer_dependencies = generate_customer_dependencies(customer_name, mom_text)
    print(f"[SOW] [12/15] ✓ Customer Dependencies generated ({len(customer_dependencies)} chars)")

    print(f"[SOW] [13/15] Generating 'Assumptions'...")
    assumptions = generate_assumptions(customer_name, mom_text)
    print(f"[SOW] [13/15] ✓ Assumptions generated ({len(assumptions)} chars)")

    print(f"[SOW] [14/15] Generating 'Exclusions'...")
    exclusions = generate_exclusions(customer_name, mom_text)
    print(f"[SOW] [14/15] ✓ Exclusions generated ({len(exclusions)} chars)")

    # Build real AWS Pricing Calculator estimate via MCP
    print(f"[SOW] [15/15] Building AWS Pricing Calculator estimate...")
    pricing             = build_pricing_estimate(customer_name, mom_text, key_highlights)
    print(f"[SOW] [15/15] ✓ Pricing estimate built (URL: {pricing.get('url', 'N/A')})")

    # Pass cost_summary to tco_assumptions so Bedrock generates service-specific assumptions
    cost_summary        = pricing.get("cost_summary", "")
    print(f"[SOW] [BONUS] Generating 'TCO Assumptions' based on pricing...")
    tco_assumptions     = generate_tco_assumptions(customer_name, mom_text, cost_summary)
    print(f"[SOW] [BONUS] ✓ TCO Assumptions generated ({len(tco_assumptions)} chars)")

    print(f"[SOW] ───────────────────────────────────────────────────────────")
    print(f"[SOW] All Bedrock sections generated. Building DOCX document...")
    print(f"[SOW] ───────────────────────────────────────────────────────────")

    doc = Document()

    # ── Set document-wide font defaults ──
    _set_document_defaults(doc)

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin   = Cm(1.27)
        section.right_margin  = Cm(1.27)
        section.footer_distance = Cm(1.27)

    # ── Header and Footer ──
    _add_header_footer(doc, customer_name)

    # ── Cover Page ────────────────────────────────────────────────────────
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Proposal for AWS Cloud Services")
    run.bold = False
    run.font.name = FONT_NAME
    run.font.size = Pt(22)
    run.font.color.rgb = DARK

    doc.add_paragraph()

    for_p = doc.add_paragraph()
    for_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for_run = for_p.add_run("For")
    for_run.font.name = FONT_NAME
    for_run.font.size = Pt(12)
    for_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    # ── Client Logo (centered) ──
    if client_logo_path and os.path.exists(client_logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = logo_p.add_run()
        run_logo.add_picture(client_logo_path, width=Inches(2.5))
    else:
        # If no logo uploaded, show customer name in its place
        cust_p = doc.add_paragraph()
        cust_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cust_p.add_run(customer_name)
        run2.bold = True
        run2.font.name = FONT_NAME
        run2.font.size = Pt(20)
        run2.font.color.rgb = ORANGE

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # ── Document History heading ──
    dh_heading = doc.add_paragraph()
    run_dh = dh_heading.add_run("Document History")
    run_dh.bold = True
    run_dh.font.name = "Calibri (Body)"
    run_dh.font.size = Pt(18)
    run_dh.font.color.rgb = DARK
    # Small caps effect via caps
    run_dh.font.small_caps = True

    # Document History table
    hist_table = doc.add_table(rows=2, cols=3)
    hist_table.style = "Table Grid"
    headers = ["Version", "Date", "Submitted by"]
    for i, h in enumerate(headers):
        cell = hist_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        _set_cell_bg(cell, "1F2763")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hist_table.rows[1].cells[0].text = "1.0"
    hist_table.rows[1].cells[1].text = doc_date if doc_date else "Date"
    hist_table.rows[1].cells[2].text = submitted_by if submitted_by else "Operisoft Technologies"
    # Set font for data row
    for i in range(3):
        for para in hist_table.rows[1].cells[i].paragraphs:
            for run in para.runs:
                run.font.name = "Calibri (Body)"
                run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Operisoft Logo (static, centered at bottom of cover page) ──
    operisoft_logo_path = os.path.join(os.path.dirname(__file__), "OperisoftLogo.png")
    if os.path.exists(operisoft_logo_path):
        opi_logo_p = doc.add_paragraph()
        opi_logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_opi = opi_logo_p.add_run()
        run_opi.add_picture(operisoft_logo_path, width=Inches(2.0))
    else:
        # Fallback: styled text if logo image not available
        opi_p = doc.add_paragraph()
        opi_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_opi_text = opi_p.add_run("Operisoft.")
        run_opi_text.bold = True
        run_opi_text.font.name = FONT_NAME
        run_opi_text.font.size = Pt(26)
        run_opi_text.font.color.rgb = DARK

    doc.add_page_break()

    # ── Table of Contents (Page 2) ───────────────────────────────────────
    _build_toc(doc, project_type, include_post_deploy, include_landing_zone,
               include_control_tower, include_lz_arch, include_paloalto,
               include_mgn, include_testing_monitoring, include_monitoring, include_dr)

    doc.add_page_break()

    # ── 1. Executive Summary ──────────────────────────────────────────────
    _heading(doc, "1. Executive Summary", 1, DARK)

    _heading(doc, "1.1 Confidentiality Notice", 2, ORANGE)
    _para(doc, CONFIDENTIALITY_TEXT)
    _section_divider(doc)

    _heading(doc, "1.2 About Operisoft", 2, ORANGE)
    for para in ABOUT_OPERISOFT.split("\n\n"):
        _para(doc, para.strip())
    _section_divider(doc)

    _heading(doc, "1.3 About Customer", 2, ORANGE)
    _para(doc, about_customer)
    _section_divider(doc)

    # ── 2. Project Details ────────────────────────────────────────────────
    _heading(doc, "2. Project Details", 1, DARK)

    _heading(doc, "2.1 Project Objectives and Overview", 2, ORANGE)
    _render_ai_content(doc, proj_objectives)
    _section_divider(doc)

    _heading(doc, "2.2 Understanding of Current Landscape", 2, ORANGE)
    _render_ai_content(doc, current_landscape)
    _section_divider(doc)

    # ── 3. Proposed Solution ──────────────────────────────────────────────
    _heading(doc, "3. Proposed Solution", 1, DARK)

    _heading(doc, "3.1 High Level Scope", 2, ORANGE)
    _render_ai_content(doc, high_level_scope)
    _section_divider(doc)

    _heading(doc, "3.2 Scope of Work", 2, ORANGE)
    _render_ai_content(doc, scope_of_work)
    _section_divider(doc)

    _heading(doc, "3.3 Key Highlights of the Solution", 2, ORANGE)
    _render_ai_content(doc, key_highlights)
    _section_divider(doc)

    _heading(doc, "3.4 Proposed Solution Diagram", 2, ORANGE)
    _para(doc, "[Proposed Solution Architecture Diagram – Please insert diagram here]", italic=True, color=RGBColor(0x99, 0x99, 0x99))
    _section_divider(doc)

    # 3.5 Project Deliverables (Bedrock-generated)
    _heading(doc, "3.5 Project Deliverables", 2, ORANGE)
    _render_ai_content(doc, project_deliverables)
    _section_divider(doc)

    # 3.6 heading is conditional on project type
    if project_type.strip().upper() == "POC":
        ac_heading = "3.6 POC Acceptance Criteria"
    else:
        ac_heading = "3.6 Project Acceptance Criteria"
    _heading(doc, ac_heading, 2, ORANGE)
    _render_ai_content(doc, acceptance_criteria)
    _section_divider(doc)

    # Optional: 4. Post-Deployment Testing and Acceptance
    if include_post_deploy:
        _heading(doc, "4. Post-Deployment Testing and Acceptance", 1, DARK)
        for b in POST_DEPLOYMENT_BULLETS:
            _bullet(doc, b)
        _section_divider(doc)

    # Optional sections 3.7, 3.8, 3.9, 3.10
    sec_num = 7
    if include_landing_zone:
        _heading(doc, "3.7 Implementation of Landing Zone", 2, ORANGE)
        _para(doc, "In this solution, we propose setting up a Landing zone on the newly created AWS account setup. We are proposing the following Landing Zone architecture:")
        doc.add_paragraph()
        lz_table = doc.add_table(rows=1, cols=3)
        lz_table.style = "Table Grid"
        for i, h in enumerate(["Organization Root", "Organization Unit", "Account"]):
            cell = lz_table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, "1F2763")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        ou_data = [
            ("Organizational Root Account", "Master", "Management"),
            ("", "Security (OU)", "Security"),
            ("", "", "Log Archive"),
            ("", "Infra (OU)", "Network"),
            ("", "", "Observability"),
            ("", "", "Shared Services"),
            ("", "Sandbox (OU)", "Account for UAT and Dev application deployment"),
            ("", "Workload (OU)", "Account for Production application deployment"),
        ]
        for row_data in ou_data:
            row = lz_table.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = val
        _section_divider(doc)
        _para(doc, "There will be a single master account for the entire organization. Under the organization's root, there will be a central security OU and separate OUs for each business unit/LOB.")
        _para(doc, "Mandatory Guardrails:", bold=True)
        for g in LANDING_ZONE_MANDATORY_GUARDRAILS:
            _bullet(doc, g)
        _para(doc, "Recommended Guardrails:", bold=True)
        for g in LANDING_ZONE_RECOMMENDED_GUARDRAILS:
            _bullet(doc, g)
        _section_divider(doc)
        sec_num += 1

    if include_control_tower:
        _heading(doc, "3.8 Configuration of Control Tower Setup", 2, ORANGE)
        ct_bullets = [
            "Enable Control Tower from AWS Console",
            "Register OU to Control Tower using the console",
            "Creation of a new Account in the Organization using Automation",
            "Configuration of service control policies using control tower customization automation",
            "Configuration of security Hub using CloudFormation",
            "Configuration of GuardDuty using CloudFormation",
        ]
        for b in ct_bullets:
            _bullet(doc, b)
        _section_divider(doc)

    if include_lz_arch:
        _heading(doc, "3.9 AWS Landing Zone Architecture", 2, ORANGE)
        lza_bullets = [
            "End user account provisioning through AWS Service Catalog (Account factory)",
            "Centralized monitoring and notifications for each OU using Amazon CloudWatch and Amazon SNS",
            "Centralized logging and immutable archive for all Cloud infrastructure and application logs",
            "Workload OU and Accounts to be defined for separate Business Units, applications, or environments",
            "Isolation of resources using separate accounts for similar groups of workloads",
            "Automation for compliance management and security posture management",
            "Automation for Organization policies like SCP, backup, and Tagging policies",
            "Automatic Tagging of resources for Billing segmentation for CloudFinOps",
        ]
        for b in lza_bullets:
            _bullet(doc, b)
        _section_divider(doc)

    if include_paloalto:
        _heading(doc, "3.10 PaloAlto Next Generation Firewall", 2, ORANGE)
        _para(doc, "We propose PaloAlto Next Generation Firewall with advanced security features like IPS, URL filtering, Bot Protection, Sandboxing which will inspect all the incoming and outgoing traffic.")
        _para(doc, "Key VM-Series Features and Capabilities:", bold=True)
        for b in PALOALTO_BULLETS:
            _bullet(doc, b)
        _section_divider(doc)

    # ── Migration Approach ───────────────────────────────────────────
    if include_mgn or include_testing_monitoring:
        _heading(doc, "4. Migration Approach", 1, DARK)

        if include_mgn:
            _heading(doc, "4.1 Migration of Data Using MGN", 2, ORANGE)
            for b in MGN_BULLETS:
                _bullet(doc, b)
            _section_divider(doc)

        if include_testing_monitoring:
            _heading(doc, "4.2 Testing and Monitoring", 2, ORANGE)
            for t in TESTING_MONITORING_TEXT:
                _bullet(doc, t)
            _section_divider(doc)

    # ── Monitoring ───────────────────────────────────────────────────────
    if include_monitoring:
        _heading(doc, "5. Monitoring AWS Infrastructure", 1, DARK)
        for b in MONITORING_INFRA_BULLETS:
            _bullet(doc, b)
        _section_divider(doc)

    # ── DR ───────────────────────────────────────────────────────────────
    if include_dr:
        _heading(doc, "6. Proposed DR Approach", 1, DARK)

        _heading(doc, "6.1 Requirement of DR", 2, ORANGE)
        dr_rows = [
            ("Identify Critical Systems", str(dr_reqs.get("critical_systems", "All production systems"))),
            ("Primary Site Location",     str(dr_reqs.get("primary_site", "Mumbai Region"))),
            ("Type of DR",                str(dr_reqs.get("dr_type", "Passive"))),
            ("Expected RTO",              str(dr_reqs.get("rto", "30 minutes"))),
            ("Expected RPO",              str(dr_reqs.get("rpo", "15 minutes"))),
        ]
        _two_col_table(doc, dr_rows)
        _section_divider(doc)

        _heading(doc, "6.2 AWS Elastic Disaster Recovery for Replication", 2, ORANGE)
        for para in EDR_TEXT.split("\n\n"):
            _para(doc, para.strip())
        _section_divider(doc)

        _heading(doc, "6.3 AWS DRS Architecture Design", 2, ORANGE)
        _para(doc, "[AWS DRS Architecture Diagram – Please insert diagram here]", italic=True, color=RGBColor(0x99, 0x99, 0x99))
        for b in DRS_ARCH_BULLETS:
            _bullet(doc, b)
        _section_divider(doc)

        _heading(doc, "6.4 Proposed DRS Solution", 2, ORANGE)
        _render_ai_content(doc, proposed_drs)
        _section_divider(doc)

    # ── 5. Cost Estimate / Commercial Terms ──────────────────────────────
    _heading(doc, "5. Cost Estimate / Commercial Terms", 1, DARK)

    # 5.1 AWS TCO
    _heading(doc, "5.1 AWS Total Cost of Ownership [TCO]", 2, ORANGE)

    calc_url = pricing.get("url", "https://calculator.aws/pricing/2/estimate")

    link_para = doc.add_paragraph()
    link_para.paragraph_format.space_before = Pt(4)
    link_para.paragraph_format.space_after  = Pt(4)
    label_run = link_para.add_run("AWS Pricing Calculator Estimate: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    label_run.font.color.rgb = DARK
    url_run = link_para.add_run(calc_url)
    url_run.font.size = Pt(11)
    url_run.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
    url_run.underline = True

    _section_divider(doc)

    # 5.2 Operisoft Service Charges
    _heading(doc, "5.2 Operisoft Service Charges", 2, ORANGE)
    svc_table = doc.add_table(rows=3, cols=2)
    svc_table.style = "Table Grid"
    cost_rows = [
        ("Cost Component", "Cost in INR"),
        ("OTC [One Time Charges]", ""),
        ("Retainership cost [Recurring Managed services]", ""),
    ]
    for i, (k, v) in enumerate(cost_rows):
        svc_table.rows[i].cells[0].text = k
        svc_table.rows[i].cells[1].text = v
        if i == 0:
            svc_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            svc_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    _para(doc, "GST Extra as applicable", italic=True)
    _section_divider(doc)

    # 5.3 Estimated Timeline
    _heading(doc, "5.3 Estimated Timeline", 2, ORANGE)
    _bullet(doc, "4 Weeks Development.")
    _bullet(doc, "1 Week UAT.")
    _section_divider(doc)

    # 5.4 Assumptions Taken for Cost Calculation (Bedrock — driven by TCO services)
    _heading(doc, "5.4 Assumptions Taken for Cost Calculation", 2, ORANGE)
    _render_ai_content(doc, tco_assumptions)
    _section_divider(doc)

    # ── 6. Customer Obligations and Engagement Terms ────────────────────
    _heading(doc, "6. Customer Obligations and Engagement Terms", 1, DARK)
    _render_ai_content(doc, customer_obligations)
    _section_divider(doc)

    # ── 7. Customer Dependencies ──────────────────────────────────────────
    _heading(doc, "7. Customer Dependencies", 1, DARK)
    _render_ai_content(doc, customer_dependencies)
    _section_divider(doc)

    # ── 8. Assumptions ────────────────────────────────────────────────────
    _heading(doc, "8. Assumptions", 1, DARK)
    _render_ai_content(doc, assumptions)
    _section_divider(doc)

    # ── 9. Exclusions ─────────────────────────────────────────────────────
    _heading(doc, "9. Exclusions", 1, DARK)
    _render_ai_content(doc, exclusions)
    _section_divider(doc)

    # ── 10. Commercial Terms — always starts on a new page ──────────────
    doc.add_page_break()
    _heading(doc, "10  Commercial Terms and Conditions", 1, DARK)

    # Helper for roman numeral items (i, ii, iii, etc.)
    def _roman_item(doc, label, bold_text, normal_text=""):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Cm(-1.5)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run_label = p.add_run(f"{label}.\t")
        run_label.font.name = "Calibri (Body)"
        run_label.font.size = Pt(11)
        run_bold = p.add_run(bold_text)
        run_bold.bold = True
        run_bold.font.name = "Calibri (Body)"
        run_bold.font.size = Pt(11)
        if normal_text:
            run_norm = p.add_run(normal_text)
            run_norm.font.name = "Calibri (Body)"
            run_norm.font.size = Pt(11)
        return p

    def _sub_item(doc, label, text, bold_part=""):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(3.0)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.paragraph_format.space_after = Pt(2)
        run_label = p.add_run(f"{label}.\t")
        run_label.font.name = "Calibri (Body)"
        run_label.font.size = Pt(11)
        if bold_part:
            run_b = p.add_run(bold_part)
            run_b.bold = True
            run_b.font.name = "Calibri (Body)"
            run_b.font.size = Pt(11)
        run_text = p.add_run(text)
        run_text.font.name = "Calibri (Body)"
        run_text.font.size = Pt(11)
        return p

    def _indented_line(doc, text, indent=3.0, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = "Calibri (Body)"
        run.font.size = Pt(11)
        run.bold = bold
        return p

    # i. Offer Validity
    _roman_item(doc, "i", "Offer Validity: ", "Upto 30 days.")

    # ii. Payment Term
    _roman_item(doc, "ii", "Payment Term:")
    _sub_item(doc, "a", "Implementation: 100% advance unless explicitly mentioned.", "")
    _sub_item(doc, "b", "Managed Services/AWS Bills: 15 days after Invoice generation.", "")

    # iii. Taxes
    _roman_item(doc, "iii", "Taxes: ", "Extra @ 18% IGST/CGST+SGST in case of INR billing")

    # iv. SAC/HSN
    _roman_item(doc, "iv", "SAC/HSN:")
    _sub_item(doc, "a", " Cloud consumption & services", "998315")
    _sub_item(doc, "b", " Software", "997331")

    # v. Billing Cycle
    _roman_item(doc, "v", "Billing Cycle:")
    p_bill = doc.add_paragraph(style="List Bullet")
    p_bill.paragraph_format.left_indent = Cm(3.0)
    p_bill.paragraph_format.space_after = Pt(4)
    run_bill = p_bill.add_run("Monthly invoices are generated from 5th to 15th of the next month. And must be paid within ")
    run_bill.font.name = "Calibri (Body)"
    run_bill.font.size = Pt(11)
    run_bill_bold = p_bill.add_run("15 days")
    run_bill_bold.bold = True
    run_bill_bold.font.name = "Calibri (Body)"
    run_bill_bold.font.size = Pt(11)
    run_bill2 = p_bill.add_run(" to avoid stopping the services.")
    run_bill2.font.name = "Calibri (Body)"
    run_bill2.font.size = Pt(11)

    # vi. Order to be placed on
    _roman_item(doc, "vi", "Order to be placed on:")
    _indented_line(doc, "OPERISOFT TECHNOLOGIES PRIVATE LIMITED", indent=3.0)
    _indented_line(doc, "Office No 301, Plot No 273, Zion Complex Building", indent=3.0)
    _indented_line(doc, "Sector-10, Kharghar, Navi Mumbai \u2013 410210", indent=3.0)

    # vii. Mode of Payment
    _roman_item(doc, "vii", "Mode of Payment")
    doc.add_paragraph()
    _indented_line(doc, "By Wire Transfer/NEFT:", indent=2.0, bold=True)

    # Payment details as tabbed lines (not table)
    pay_data = [
        ("Beneficiary Name", "OPERISOFT TECHNOLOGIES PVT LTD"),
        ("Bank Name & Add", "ICICI BANK, Vashi Branch, Navi Mumbai"),
        ("Bank A/c No.", "015105019703"),
        ("Account Type", "Current Account"),
        ("IFSC/IMPS/RTGS CODE", "ICIC0000151"),
        ("SWIFT CODE", "ICICINBBCTS"),
    ]
    for label, value in pay_data:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(3.0)
        p.paragraph_format.space_after = Pt(2)
        run_l = p.add_run(f"{label}\t")
        run_l.bold = True
        run_l.font.name = "Calibri (Body)"
        run_l.font.size = Pt(11)
        run_v = p.add_run(value)
        run_v.font.name = "Calibri (Body)"
        run_v.font.size = Pt(11)
        # Add a right-aligned tab stop for alignment
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:pos'), '5670')  # ~10cm from left margin
        tabs.append(tab)
        pPr.append(tabs)

    doc.add_paragraph()
    p_gst = doc.add_paragraph()
    p_gst.paragraph_format.left_indent = Cm(2.0)
    run_gst = p_gst.add_run("Our (GST Number) GSTIN # ")
    run_gst.font.name = "Calibri (Body)"
    run_gst.font.size = Pt(11)
    run_gst_val = p_gst.add_run("27AADCO8045M1ZZ")
    run_gst_val.bold = True
    run_gst_val.font.name = "Calibri (Body)"
    run_gst_val.font.size = Pt(11)

    doc.add_paragraph()

    # DISCLAIMER
    p_disc_heading = doc.add_paragraph()
    p_disc_heading.paragraph_format.space_before = Pt(12)
    run_disc = p_disc_heading.add_run("DISCLAIMER:")
    run_disc.bold = True
    run_disc.font.name = "Calibri (Body)"
    run_disc.font.size = Pt(11)

    disc_items = [
        "Training or anything else which is not mentioned in above document is out of scope.",
        "For any feature/functionalities to be added which is not in scope will be charged extra as per mutual agreement.",
        "Using the documents or data for any purpose other than this, should be taken in writing by Operisoft.",
    ]
    for idx, item in enumerate(disc_items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Cm(-1.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run_label = p.add_run(f"{['i', 'ii', 'iii'][idx]}.\t")
        run_label.font.name = "Calibri (Body)"
        run_label.font.size = Pt(9)
        run_text = p.add_run(item)
        run_text.font.name = "Calibri (Body)"
        run_text.font.size = Pt(9)

    # ── Signature Page (last page) ────────────────────────────────────────
    doc.add_page_break()

    # "Signatures for this Agreement" banner — blue background, white text, centered
    sig_table_header = doc.add_table(rows=1, cols=1)
    sig_table_header.style = "Table Grid"
    sig_cell = sig_table_header.rows[0].cells[0]
    sig_cell.text = "Signatures for this Agreement"
    sig_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_cell.paragraphs[0].runs[0].font.name = "Calibri (Body)"
    sig_cell.paragraphs[0].runs[0].font.size = Pt(14)
    sig_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_bg(sig_cell, "1F2763")

    doc.add_paragraph()

    # IN WITNESS WHEREOF paragraph
    p_witness = doc.add_paragraph()
    p_witness.paragraph_format.space_after = Pt(12)
    run_bold = p_witness.add_run("IN WITNESS WHEREOF")
    run_bold.bold = True
    run_bold.font.name = "Calibri (Body)"
    run_bold.font.size = Pt(11)
    run_rest = p_witness.add_run(", the parties hereto each acting with proper authority, for good and valuable consideration and pursuant to the terms of the Agreement have executed this Statement of Work.")
    run_rest.font.name = "Calibri (Body)"
    run_rest.font.size = Pt(11)

    doc.add_paragraph()

    # Signature table — 2 columns, 4 rows
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"

    # Row 0: Signature lines + company names
    cell_left = sig_table.rows[0].cells[0]
    cell_left.paragraphs[0].clear()
    p_l = cell_left.paragraphs[0]
    p_l.paragraph_format.space_before = Pt(40)
    run_line_l = p_l.add_run("________________________\n")
    run_line_l.font.name = "Calibri (Body)"
    run_line_l.font.size = Pt(11)
    run_company = p_l.add_run("(Operisoft Technologies Pvt Ltd)")
    run_company.bold = True
    run_company.font.name = "Calibri (Body)"
    run_company.font.size = Pt(11)

    cell_right = sig_table.rows[0].cells[1]
    cell_right.paragraphs[0].clear()
    p_r = cell_right.paragraphs[0]
    p_r.paragraph_format.space_before = Pt(40)
    run_line_r = p_r.add_run("________________________\n")
    run_line_r.font.name = "Calibri (Body)"
    run_line_r.font.size = Pt(11)
    run_cust = p_r.add_run(f"({customer_name})")
    run_cust.bold = False
    run_cust.font.name = "Calibri (Body)"
    run_cust.font.size = Pt(11)

    # Row 1: Name
    sig_table.rows[1].cells[0].text = "Name:"
    sig_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    sig_table.rows[1].cells[0].paragraphs[0].runs[0].font.name = "Calibri (Body)"
    sig_table.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    sig_table.rows[1].cells[1].text = "Name:"
    sig_table.rows[1].cells[1].paragraphs[0].runs[0].bold = True
    sig_table.rows[1].cells[1].paragraphs[0].runs[0].font.name = "Calibri (Body)"
    sig_table.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(11)

    # Row 2: Title
    sig_table.rows[2].cells[0].text = "Title:"
    sig_table.rows[2].cells[0].paragraphs[0].runs[0].bold = True
    sig_table.rows[2].cells[0].paragraphs[0].runs[0].font.name = "Calibri (Body)"
    sig_table.rows[2].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    sig_table.rows[2].cells[1].text = "Title:"
    sig_table.rows[2].cells[1].paragraphs[0].runs[0].bold = True
    sig_table.rows[2].cells[1].paragraphs[0].runs[0].font.name = "Calibri (Body)"
    sig_table.rows[2].cells[1].paragraphs[0].runs[0].font.size = Pt(11)

    # Row 3: Signature Date
    sig_table.rows[3].cells[0].text = "Signature Date:"
    sig_table.rows[3].cells[0].paragraphs[0].runs[0].bold = True
    sig_table.rows[3].cells[0].paragraphs[0].runs[0].font.name = "Calibri (Body)"
    sig_table.rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    sig_table.rows[3].cells[1].text = "Signature Date:"
    sig_table.rows[3].cells[1].paragraphs[0].runs[0].bold = True
    sig_table.rows[3].cells[1].paragraphs[0].runs[0].font.name = "Calibri (Body)"
    sig_table.rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # ── Save file ──────────────────────────────────────────────────────────
    output_dir = tempfile.gettempdir()
    filename = f"SOW_{customer_name.replace(' ', '_')}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(f"[SOW] ═══════════════════════════════════════════════════════════")
    print(f"[SOW] ✓ SOW document saved: {output_path}")
    print(f"[SOW] ✓ Generation complete for: {customer_name}")
    print(f"[SOW] ═══════════════════════════════════════════════════════════\n")
    return output_path
